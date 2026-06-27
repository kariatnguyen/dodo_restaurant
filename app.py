import re

from flask import Flask, render_template, request, redirect, url_for, session
import mysql.connector
from flask import Flask, render_template
import os

app = Flask(__name__)
app.secret_key = "DODO_SECRET_KEY"


# =========================================================
# KẾT NỐI MYSQL
# =========================================================
class MySQLConnection:

    def __init__(self):
        self.host = "mysql.railway.internal"
        self.user = "root"
        self.password = "dioCSfUNCDGREmhJKxFCtzQngIkEnAUU"
        self.database = "railway"

    def connect(self):
        try:
            conn = mysql.connector.connect(
                host=self.host,
                user=self.user,
                password=self.password,
                database=self.database
            )
            return conn

        except mysql.connector.Error as err:
            print("Connection Error:", err)
            return None


# =========================================================
# LOGIN
# =========================================================

from flask import render_template, request, redirect, url_for, session

@app.route("/login", methods=["GET", "POST"])
def login():

    # Hiển thị trang đăng nhập
    if request.method == "GET":
        return render_template("login.html")

    # Lấy dữ liệu từ form
    username = request.form.get("username", "").strip()
    password = request.form.get("password", "").strip()

    # Kiểm tra dữ liệu
    if username == "" or password == "":
        return render_template(
            "login.html",
            error="Vui lòng nhập đầy đủ tài khoản và mật khẩu."
        )

    # Kết nối MySQL
    conn = MySQLConnection().connect()

    if conn is None:
        return render_template(
            "login.html",
            error="Không thể kết nối cơ sở dữ liệu."
        )

    try:

        cursor = conn.cursor(dictionary=True)

        sql = """
            SELECT
                e.employee_id,
                e.name,
                e.phone,
                e.username,
                e.password,
                e.role_id,
                e.status,
                r.role_name
            FROM employees e
            INNER JOIN roles r
                ON e.role_id = r.role_id
            WHERE e.username = %s
              AND e.password = %s
            LIMIT 1
        """

        cursor.execute(sql, (username, password))

        account = cursor.fetchone()

        # Sai tài khoản hoặc mật khẩu
        if account is None:
            return render_template(
                "login.html",
                error="Sai tên đăng nhập hoặc mật khẩu."
            )

        # Kiểm tra trạng thái
        if account["status"] != "active":
            return render_template(
                "login.html",
                error="Tài khoản đã bị khóa."
            )

        # Lưu Session
        session["employee_id"] = account["employee_id"]
        session["username"] = account["username"]
        session["name"] = account["name"]
        session["role_id"] = account["role_id"]
        session["role_name"] = account["role_name"]

        # Phân quyền
        if account["role_id"] == 1:
            print("Admin logged in")
            return redirect(url_for("admin_dashboard"))

        elif account["role_id"] == 2:
            return redirect(url_for("kitchen_dashboard"))

        elif account["role_id"] == 3:
            return redirect(url_for("staff_dashboard"))

        else:
            session.clear()

            return render_template(
                "login.html",
                error="Quyền truy cập không hợp lệ."
            )

    except mysql.connector.Error as err:

        print(err)

        return render_template(
            "login.html",
            error="Đã xảy ra lỗi trong quá trình đăng nhập."
        )

    finally:

        cursor.close()
        conn.close()

# =========================================================
# ADMIN
# =========================================================

@app.route("/admin")
@app.route("/admin/dashboard")
def admin_dashboard():

    # Chưa đăng nhập
    if "employee_id" not in session:
        return redirect(url_for("login"))

    # Không phải Admin
    if session.get("role_id") != 1:
        return redirect(url_for("login"))

    # Thông tin người đăng nhập
    employee = {
        "employee_id": session.get("employee_id"),
        "username": session.get("name"),
        "role_id": session.get("role_id"),
        "role_name": session.get("role_name")
    }

    return render_template(
        "admin/index.html",
        employee=employee
    )


@app.route("/admin/menu")
def admin_menu():

    # Kiểm tra đăng nhập
    if "employee_id" not in session:
        return redirect(url_for("login"))

    # Kiểm tra quyền Admin
    if session.get("role_id") != 1:
        return redirect(url_for("login"))

    conn = MySQLConnection().connect()

    if conn is None:
        return render_template(
            "admin/menu.html",
            categories=[],
            menu_items=[],
            error="Không thể kết nối cơ sở dữ liệu."
        )

    try:

        cursor = conn.cursor(dictionary=True)

        # =========================
        # Lấy danh sách danh mục
        # =========================
        cursor.execute("""
            SELECT
                category_id,
                name
            FROM menu_categories
            ORDER BY category_id
        """)

        categories = cursor.fetchall()

        # =========================
        # Lấy danh sách món ăn
        # =========================
        cursor.execute("""
            SELECT
                mi.item_id,
                mi.name AS item_name,
                mi.price,
                mi.status,
                mi.img,
                mc.category_id,
                mc.name AS category_name
            FROM menu_items mi
            INNER JOIN menu_categories mc
                ON mi.category_id = mc.category_id
            ORDER BY
                mc.category_id,
                mi.name
        """)

        menu_items = cursor.fetchall()

        return render_template(
            "admin/menu.html",
            categories=categories,
            menu_items=menu_items
        )

    except mysql.connector.Error as err:

        print(err)

        return render_template(
            "admin/menu.html",
            categories=[],
            menu_items=[],
            error="Có lỗi xảy ra khi tải dữ liệu."
        )

    finally:

        cursor.close()
        conn.close()

# Thêm danh mục
@app.route("/admin/menu/add_category", methods=["POST"])
def add_category():
    # Implementation for adding category
    name = request.form['name']
    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("INSERT INTO menu_categories (name) VALUES (%s)", (name,))
        conn.commit()
    except mysql.connector.Error as err:
        print(err)
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for("admin_menu"))

# Sửa danh mục
@app.route("/admin/menu/update_category", methods=["POST"])
def update_category():
    category_id = request.form['category_id']
    name = request.form['name']
    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("UPDATE menu_categories SET name = %s WHERE category_id = %s", (name, category_id))
        conn.commit()
    except mysql.connector.Error as err:
        print(err)
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for("admin_menu"))

# Xóa danh mục
# <button class="action-btn delete" onclick="window.location.href='/admin/menu/delete_category?category_id={{ category.category_id }}'"> Xóa </button>
@app.route("/admin/menu/delete_category", methods=["GET"])
def delete_category():
    category_id = request.args.get("category_id")
    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("""DELETE FROM menu_items WHERE category_id = %s""", (category_id,))
        cursor.execute("DELETE FROM menu_categories WHERE category_id = %s", (category_id,))
        conn.commit()
    except mysql.connector.Error as err:
        print(err)
    finally:
        cursor.close()
        conn.close()
    return redirect(url_for("admin_menu"))

# Thêm sản phẩm
@app.route("/admin/menu/add_product", methods=["POST"])
def add_product():

    name = request.form['name']
    price = request.form['price']
    category_id = request.form['category_id']
    img = request.form['img']
    status = request.form['status']

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            INSERT INTO menu_items (category_id, name, price, img, status)
            VALUES (%s, %s, %s, %s, %s)
        """, (category_id, name, price, img, status))
        print("Product added successfully.")
        conn.commit()

    except Exception as e:
        print("ERROR INSERT:", e)

    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('admin_menu', request='menu'))

# Xóa sản phẩm
@app.route("/admin/menu/delete_product/<int:item_id>")
def delete_menu_product(item_id):

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            DELETE FROM menu_items
            WHERE item_id = %s
        """, (item_id,))

        conn.commit()

    except Exception as e:
        print("DELETE ERROR:", e)

    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('admin_menu', request='menu'))

# Sửa sản phẩm
@app.route("/admin/menu/update_product", methods=["POST"])
def update_product():

    item_id = request.form['item_id']
    name = request.form['name']
    price = request.form['price']
    category_id = request.form['category_id']
    img = request.form['img']
    status = request.form['status']

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            UPDATE menu_items
            SET category_id = %s, name = %s, price = %s, img = %s, status = %s
            WHERE item_id = %s
        """, (category_id, name, price, img, status, item_id))

        conn.commit()

    except Exception as e:
        print("UPDATE ERROR:", e)

    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('admin_menu', request='menu'))

# Thay đổi trạng thái sản phẩm
@app.route("/admin/menu/toggle_status/<int:item_id>/<status>")
def toggle_status(item_id, status):

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            UPDATE menu_items
            SET status = %s
            WHERE item_id = %s
        """, (status, item_id))

        conn.commit()

    except Exception as e:
        print("TOGGLE ERROR:", e)

    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('admin_menu', request='menu'))

@app.route("/admin/staff")
def admin_staff():

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""
        SELECT 
            e.employee_id,
            e.name,
            e.username,
            e.status,
            r.role_name,
            e.created_at
        FROM employees e
        JOIN roles r ON e.role_id = r.role_id
    """)

    staff = cursor.fetchall()

    cursor.execute("SELECT * FROM roles")
    roles = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        "admin/staff.html",
        staff=staff,
        roles=roles
    )

@app.route("/admin/table-management")
def table_dashboard():

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""
        SELECT t.table_id, t.table_name, o.order_id, o.order_time,
               COALESCE(SUM(p.amount), 0) AS total_revenue
        FROM tables t
        JOIN orders o ON o.table_id = t.table_id
        LEFT JOIN payments p ON p.order_id = o.order_id
        WHERE o.status = 'active'
        GROUP BY t.table_id, o.order_id, o.order_time
    """)
    active_tables = cursor.fetchall()

    cursor.execute("""
        SELECT r.reservation_id, r.table_id, c.name AS customer_name,
               r.reservation_time, r.status
        FROM reservations r
        JOIN customers c ON c.customer_id = r.customer_id
        WHERE r.status = 'booked'
        ORDER BY r.reservation_time ASC
    """)
    bookings = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        "admin/table_management.html",
        active_tables=active_tables,
        bookings=bookings
    )

# thêm nhân viên
@app.route("/admin/staff/add", methods=["POST"])
def add_staff():

    name = request.form["name"]
    username = request.form["username"]
    role_id = request.form["role_id"]
    password = request.form["password"]
    phone = request.form["phone"]

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    cursor.execute("""
        INSERT INTO employees (name, phone, username, password, role_id, status)
        VALUES (%s, %s, %s, %s, %s, 'active')
    """, (name, phone, username, password, role_id))

    conn.commit()

    cursor.close()
    conn.close()

    return redirect(url_for("admin_staff"))

# Mở khóa nhân viên
@app.route("/admin/staff/lock/<int:id>")
def lock_staff(id):

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    cursor.execute("""
        UPDATE employees
        SET status = IF(status='active','locked','active')
        WHERE employee_id=%s
    """, (id,))

    conn.commit()
    cursor.close()
    conn.close()

    return redirect(url_for("admin_staff"))

# Sửa thông tin nhân viên
@app.route("/admin/staff/save", methods=["POST"])
def save_staff():
    employee_id = request.form.get("employee_id")

    name = request.form["name"]
    username = request.form["username"]
    phone = request.form["phone"]
    password = request.form["password"]
    role_id = request.form["role_id"]

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    if employee_id:  
        # UPDATE
        cursor.execute("""
            UPDATE employees
            SET name=%s, username=%s, phone=%s, role_id=%s
            WHERE employee_id=%s
        """, (name, username, phone, role_id, employee_id))
    else:
        # INSERT
        cursor.execute("""
            INSERT INTO employees(name, username, phone, password, role_id, status)
            VALUES (%s,%s,%s,%s,%s,'ACTIVE')
        """, (name, username, phone, password, role_id))

    conn.commit()
    cursor.close()
    conn.close()

    return redirect(url_for("admin_staff"))

from flask_mail import Mail, Message

app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'kariatnguyen@gmail.com'
app.config['MAIL_PASSWORD'] = 'vcij rkzx hfnh imbp'

mail = Mail(app)

# Gửi mail xác nhận hủy đặt bàn
@app.route("/admin/reservation/cancel/<int:id>")
def cancel_reservation(id):

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    # 1. lấy thông tin booking + email khách
    cursor.execute("""
        SELECT r.reservation_id, c.email, c.name, r.table_id, r.reservation_time
        FROM reservations r
        JOIN customers c ON c.customer_id = r.customer_id
        WHERE r.reservation_id = %s
    """, (id,))

    booking = cursor.fetchone()

    if not booking:
        return "Not found"

    # 2. update trạng thái
    cursor.execute("""
        UPDATE reservations
        SET status = 'cancelled'
        WHERE reservation_id = %s
    """, (id,))

    conn.commit()

    # 3. gửi email
    msg = Message(
        subject="Hủy đặt bàn tại nhà hàng",
        sender=app.config['MAIL_USERNAME'],
        recipients=[booking['email']]
    )

    msg.body = f"""
Xin chào {booking['name']},

Đặt bàn của bạn tại nhà hàng đã bị hủy:

- Bàn: {booking['table_id']}
- Thời gian: {booking['reservation_time']}

Nếu có thắc mắc, vui lòng liên hệ nhà hàng.

Xin cảm ơn!
"""

    mail.send(msg)

    cursor.close()
    conn.close()

    return redirect(url_for("table_dashboard"))

# Quản lý kho
@app.route("/admin/warehouse")
def warehouse():

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    # danh sách nguyên liệu
    cursor.execute("""
        SELECT * FROM ingredients
    """)
    ingredients = cursor.fetchall()

    # tổng nguyên liệu
    total_items = len(ingredients)

    # nguyên liệu sắp hết (ví dụ < 10)
    low_stock = sum(1 for i in ingredients if i['stock'] < 10)

    # tổng giá trị kho
    inventory_value = sum(i['stock'] * i['price'] for i in ingredients)

    cursor.close()
    conn.close()

    return render_template(
        "admin/warehouse.html",
        ingredients=ingredients,
        total_items=total_items,
        low_stock=low_stock,
        inventory_value=inventory_value
    )

# Nhập kho theo id nguyên liệu
@app.route("/admin/warehouse/import", methods=["POST"])
def import_ingredient():

    ingredient_id = request.form['ingredient_id']
    qty = float(request.form['qty'])
    note = request.form['note']

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    try:
        # 1. tăng tồn kho
        cursor.execute("""
            UPDATE ingredients
            SET stock = stock + %s
            WHERE ingredient_id = %s
        """, (qty, ingredient_id))

        # 2. log lại lịch sử
        cursor.execute("""
            INSERT INTO inventory_logs (ingredient_id, change_qty, type, created_at)
            VALUES (%s, %s, 'IMPORT', NOW())
        """, (ingredient_id, qty))

        conn.commit()

    except Exception as e:
        print("IMPORT ERROR:", e)

    finally:
        cursor.close()
        conn.close()

    return redirect(url_for("warehouse"))

# Thêm nguyên liệu mới
@app.route("/admin/warehouse/add_ingredient", methods=["POST"])
def add_ingredient():

    name = request.form['name']
    stock = request.form['stock']
    unit = request.form['unit']
    price = request.form['price']

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            INSERT INTO ingredients (name, stock, unit, price)
            VALUES (%s, %s, %s, %s)
        """, (name, stock, unit, price))

        conn.commit()

    except Exception as e:
        print("ADD INGREDIENT ERROR:", e)

    finally:
        cursor.close()
        conn.close()

    return redirect(url_for("warehouse"))

# Xuất báo cáo kho ra file Word
from docx import Document
from flask import send_file
import io
import mysql.connector


@app.route("/admin/warehouse/export_word")
def export_warehouse_word():

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    # ====== 1. LẤY NGUYÊN LIỆU ======
    cursor.execute("SELECT * FROM ingredients")
    ingredients = cursor.fetchall()

    # ====== 2. LẤY LOG ======
    cursor.execute("""
        SELECT l.log_id, i.name, l.change_qty, l.type, l.created_at
        FROM inventory_logs l
        JOIN ingredients i ON i.ingredient_id = l.ingredient_id
        ORDER BY l.created_at DESC
    """)
    logs = cursor.fetchall()

    cursor.close()
    conn.close()

    # ====== 3. TẠO FILE WORD ======
    doc = Document()

    doc.add_heading("BÁO CÁO KHO NGUYÊN LIỆU", 0)

    # ====== THỐNG KÊ ======
    doc.add_heading("Danh sách nguyên liệu", level=1)

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Tên"
    hdr[1].text = "Tồn kho"
    hdr[2].text = "Đơn vị"
    hdr[3].text = "Giá"

    for i in ingredients:
        row = table.add_row().cells
        row[0].text = str(i["name"])
        row[1].text = str(i["stock"])
        row[2].text = str(i["unit"])
        row[3].text = str(i.get("price", 0))

    # ====== LOG ======
    doc.add_heading("Lịch sử nhập/xuất kho", level=1)

    table2 = doc.add_table(rows=1, cols=4)
    h = table2.rows[0].cells
    h[0].text = "Tên"
    h[1].text = "Số lượng"
    h[2].text = "Loại"
    h[3].text = "Thời gian"

    for l in logs:
        row = table2.add_row().cells
        row[0].text = str(l["name"])
        row[1].text = str(l["change_qty"])
        row[2].text = str(l["type"])
        row[3].text = str(l["created_at"])

    # ====== RETURN FILE ======
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="bao_cao_kho.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# =========================================================
# KITCHEN
# =========================================================

@app.route("/kitchen")
@app.route("/kitchen/dashboard")
def kitchen_dashboard():

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""
        SELECT
            oi.order_item_id,
            o.order_id,
            t.table_name,
            t.table_code,
            mi.name AS item_name,
            oi.quantity,
            oi.status,
            o.order_time
        FROM order_items oi
        JOIN orders o ON oi.order_id = o.order_id
        JOIN tables t ON o.table_id = t.table_id
        JOIN menu_items mi ON oi.item_id = mi.item_id
        WHERE oi.status IN ('waiting', 'cooking')
        ORDER BY o.order_time ASC
    """)

    items = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template("kitchen/index.html", items=items)

# hoàn thành món 
@app.route("/kitchen/done/<int:item_id>", methods=["POST"])
def done_cooking(item_id):
    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    cursor.execute("""
        UPDATE order_items
        SET status='done'
        WHERE order_item_id=%s
    """, (item_id,))

    conn.commit()
    cursor.close()
    conn.close()

    return redirect(url_for("kitchen_dashboard"))

# update trạng thái món ăn
@app.route("/kitchen/start/<int:item_id>", methods=["POST"])
def start_cooking(item_id):
    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    cursor.execute("""
        UPDATE order_items
        SET status='cooking'
        WHERE order_item_id=%s
    """, (item_id,))

    conn.commit()
    cursor.close()
    conn.close()

    return redirect(url_for("kitchen_dashboard"))
# =========================================================
# STAFF
# =========================================================

@app.route("/staff")
@app.route("/staff/dashboard")
def staff_dashboard():
    if "employee_id" not in session:
        return redirect(url_for("login"))

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    # =========================
    # 1. DANH SÁCH BÀN (ĐÚNG TRẠNG THÁI)
    # =========================
    cursor.execute("""
        SELECT 
            t.table_id,
            t.table_name,
            t.table_code,
            t.status,

            COALESCE(SUM(oi.quantity * mi.price), 0) AS revenue

        FROM tables t

        LEFT JOIN orders o 
            ON t.table_id = o.table_id 
            AND o.status IN ('cooking','waiting')

        LEFT JOIN order_items oi 
            ON o.order_id = oi.order_id

        LEFT JOIN menu_items mi 
            ON oi.item_id = mi.item_id

        GROUP BY t.table_id, t.table_name, t.status

        ORDER BY t.table_id
    """)
    tables = cursor.fetchall()

    # =========================
    # 2. ĐẶT BÀN
    # =========================
    cursor.execute("""
        SELECT 
            r.reservation_id,
            r.reservation_time,
            r.status,
            r.table_id,
            t.table_name,
            c.name AS customer_name
        FROM reservations r
        LEFT JOIN customers c ON r.customer_id = c.customer_id
        LEFT JOIN tables t ON r.table_id = t.table_id
        ORDER BY r.reservation_time ASC
    """)
    reservations = cursor.fetchall()

    # =========================
    # 3. MÓN BẾP
    # =========================
    cursor.execute("""
        SELECT 
            o.order_id,
            t.table_name,
            oi.item_id,
            mi.name AS item_name,
            oi.quantity,
            oi.status,
            o.order_time
        FROM order_items oi
        INNER JOIN orders o ON oi.order_id = o.order_id
        INNER JOIN tables t ON o.table_id = t.table_id
        INNER JOIN menu_items mi ON oi.item_id = mi.item_id
        WHERE oi.status IN ('waiting','cooking')
        ORDER BY o.order_time DESC
    """)
    kitchen_orders = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        "staff/index.html",
        tables=tables,
        reservations=reservations,
        kitchen_orders=kitchen_orders
    )

@app.route("/admin/table/book", methods=["POST"])
def book_table():

    name = request.form["customer_name"]
    phone = request.form["phone"]
    email = request.form.get("email")
    table_id = request.form["table_id"]
    time = request.form["reservation_time"]

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    # =========================
    # 1. CHECK CUSTOMER (KHÔNG TRÙNG)
    # =========================
    cursor.execute("""
        SELECT customer_id FROM customers WHERE phone=%s
    """, (phone,))
    customer = cursor.fetchone()

    if customer:
        customer_id = customer["customer_id"]
    else:
        cursor.execute("""
            INSERT INTO customers(name, phone, email)
            VALUES (%s, %s, %s)
        """, (name, phone, email))
        customer_id = cursor.lastrowid

    # =========================
    # 2. CHECK TRÙNG BÀN
    # =========================
    cursor.execute("""
        SELECT * FROM reservations
        WHERE table_id=%s
        AND reservation_time = %s
        AND status IN ('pending','booked')
    """, (table_id, time))

    conflict = cursor.fetchone()

    if conflict:
        cursor.close()
        conn.close()
        return "Bàn đã được đặt vào thời gian này", 400

    # =========================
    # 3. INSERT RESERVATION
    # =========================
    cursor.execute("""
        INSERT INTO reservations
        (customer_id, table_id, reservation_time, status, created_by)
        VALUES (%s, %s, %s, 'pending', %s)
    """, (customer_id, table_id, time, 1))

    # =========================
    # 4. UPDATE TABLE STATUS
    # =========================
    cursor.execute("""
        UPDATE tables
        SET status='reserved'
        WHERE table_id=%s
    """, (table_id,))

    conn.commit()
    cursor.close()
    conn.close()

    return redirect("/staff/dashboard")

@app.route("/staff/table/<int:table_id>/orders")
def get_table_orders(table_id):
    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""
        SELECT 
            o.order_id,
            mi.name AS item_name,
            oi.quantity,
            mi.price,
            (oi.quantity * mi.price) AS total,
            o.status
        FROM orders o
        JOIN order_items oi ON o.order_id = oi.order_id
        JOIN menu_items mi ON mi.item_id = oi.item_id
        WHERE o.table_id = %s
          AND o.status IN ('waiting','cooking','done')
    """, (table_id,))

    data = cursor.fetchall()
    cursor.close()
    conn.close()

    return {"data": data}

@app.route("/admin/reservation/confirm/<int:reservation_id>", methods=["POST"])
def confirm_reservation(reservation_id):

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    # 1. update reservation -> booked
    cursor.execute("""
        UPDATE reservations
        SET status='booked'
        WHERE reservation_id=%s
    """, (reservation_id,))

    conn.commit()
    cursor.close()
    conn.close()

    return {"status": "ok"}

import random
import string
from datetime import datetime
from flask import request, jsonify

def gen_table_code():
    return "DODO-" + "".join(random.choices(string.ascii_uppercase + string.digits, k=5))


@app.route("/staff/pay_cash", methods=["POST"])
def pay_cash():

    data = request.get_json() or {}

    table_id = data.get("table_id")
    amount = data.get("amount")

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    # 1. lấy order đang active của bàn
    cursor.execute("""
        SELECT order_id
        FROM orders
        WHERE table_id=%s AND status IN ('waiting','cooking')
        LIMIT 1
    """, (table_id,))
    order = cursor.fetchone()

    if not order:
        return jsonify({"status": "error", "msg": "no order"})

    order_id = order[0]

    # 2. insert payments (ĐÚNG TABLE BẠN ĐƯA)
    cursor.execute("""
        INSERT INTO payments
        (order_id, amount, payment_method, paid_by, paid_time)
        VALUES (%s, %s, 'cash', %s, %s)
    """, (
        order_id,
        amount,
        1,  # staff tạm
        datetime.now()
    ))

    # 3. update orders -> paid
    cursor.execute("""
        UPDATE orders
        SET status='paid'
        WHERE order_id=%s
    """, (order_id,))

    # 4. reset bàn + random table_code
    new_code = gen_table_code()

    cursor.execute("""
        UPDATE tables
        SET status='empty',
            table_code=%s
        WHERE table_id=%s
    """, (new_code, table_id))

    conn.commit()
    cursor.close()
    conn.close()

    return jsonify({"status": "ok"})

@app.route("/admin/reservation/cancel/<int:reservation_id>", methods=["POST"])
def cancelOrder_reservation(reservation_id):

    conn = MySQLConnection().connect()
    cursor = conn.cursor()

    # update trạng thái
    cursor.execute("""
        UPDATE reservations
        SET status='cancelled'
        WHERE reservation_id=%s
    """, (reservation_id,))

    conn.commit()
    cursor.close()
    conn.close()

    return {"status": "ok"}

@app.route("/select_table/<int:table_id>")
def select_table(table_id):
    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    # 1. kiểm tra trạng thái bàn
    cursor.execute("""
        SELECT status 
        FROM tables 
        WHERE table_id = %s
    """, (table_id,))
    
    table = cursor.fetchone()

    if not table:
        cursor.close()
        conn.close()
        return "Bàn không tồn tại"

    # 2. nếu đã đặt hoặc đang dùng thì chặn
    if table['status'] in ('reserved', 'occupied'):
        cursor.close()
        conn.close()
        return "Bàn đã được đặt hoặc đang sử dụng"

    # 3. cập nhật trạng thái bàn = reserved
    cursor.execute("""
        UPDATE tables
        SET status = 'reserved'
        WHERE table_id = %s
    """, (table_id,))

    conn.commit()

    cursor.close()
    conn.close()

    return redirect(url_for("staff_dashboard"))
# =========================================================
# USER
# =========================================================

@app.route("/user")
@app.route("/user/dashboard")
def user_dashboard():

    # Kiểm tra khách đã nhập mã bàn hay chưa
    if "table_id" not in session:
        return redirect(url_for("index"))

    conn = MySQLConnection().connect()

    if conn is None:
        return render_template(
            "user/index.html",
            categories=[],
            foods=[],
            error="Không thể kết nối cơ sở dữ liệu."
        )

    try:

        cursor = conn.cursor(dictionary=True)

        # Lấy thông tin bàn
        cursor.execute("""
            SELECT *
            FROM tables
            WHERE table_id = %s
        """, (session["table_id"],))

        table = cursor.fetchall()

        # Lấy danh mục món ăn
        cursor.execute("""
            SELECT *
            FROM menu_categories
        """)

        categories = cursor.fetchall()
        print("Categories:", categories)

        # Lấy danh sách món ăn đang bán Và danh mục tương ứng
        cursor.execute("""
    SELECT 
        mi.item_id,
        mi.name,
        mi.price,
        mi.img,
        mi.status,
        mc.name AS category_name
    FROM menu_items mi
    LEFT JOIN menu_categories mc
        ON mi.category_id = mc.category_id
    ORDER BY mc.name ASC
""")
        food = cursor.fetchall()
        print("Food items:", food)

        return render_template(
            "user/index.html",
            table=table,
            categories=categories,
            food=food,
        )

    except mysql.connector.Error as err:

        print("MySQL Error:", err)

        return render_template(
            "user/index.html",
            categories=[],
            foods=[],
            error="Có lỗi xảy ra khi tải dữ liệu."
        )

    finally:

        cursor.close()
        conn.close()


# =========================================================
# TRANG CHỦ
# =========================================================

@app.route("/")
def index():

    # Khách chưa có session
    if "table_id" not in session:
        return render_template("index.html")

    conn = MySQLConnection().connect()

    if conn is None:
        session.clear()
        return render_template("index.html")

    try:

        cursor = conn.cursor(dictionary=True)

        cursor.execute("""
            SELECT *
            FROM tables
            WHERE table_id=%s
        """, (session["table_id"],))

        table = cursor.fetchone()

        if table:
            return redirect(url_for("user_dashboard"))

        # Nếu bàn không còn tồn tại
        session.clear()

        return render_template("index.html")

    except mysql.connector.Error as err:

        print(err)

        session.clear()

        return render_template("index.html")

    finally:

        cursor.close()
        conn.close()


# =========================================================
# NHẬP MÃ BÀN
# =========================================================

@app.route("/join-table", methods=["POST"])
def join_table():

    table_code = request.form.get("table_code")

    conn = MySQLConnection().connect()
    cursor = conn.cursor(dictionary=True)

    # 1. tìm bàn theo mã
    cursor.execute("""
        SELECT table_id, table_name, table_code, status
        FROM tables
        WHERE table_code = %s
    """, (table_code,))

    table = cursor.fetchone()

    # 2. không tồn tại
    if not table:
        cursor.close()
        conn.close()
        return "Mã bàn không tồn tại", 404

    # 4. cập nhật trạng thái bàn
    cursor.execute("""
        UPDATE tables
        SET status = 'occupied'
        WHERE table_id = %s
    """, (table["table_id"],))

    conn.commit()
    cursor.close()
    conn.close()

    # 5. lưu session
    session["table_id"] = table["table_id"]
    session["table_name"] = table["table_name"]
    session["table_code"] = table["table_code"]

    return redirect(url_for("user_dashboard"))



@app.route("/logout")
def logout():
    session.clear()   # xoá toàn bộ session đăng nhập
    return redirect(url_for("login"))

if __name__ == '__main__':
    app.run(
        host='0.0.0.0',
        port = int(os.environ.get("PORT", 5000)),
        debug=True
    )
